from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(r"D:\Projects\50012832提交材料")
SCREEN = OUT / "文档截图"
ASSET = OUT / "文档素材"
OUT.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)

NAVY = "183B56"
BLUE = "2F75B5"
TEAL = "2E7D6B"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F4F6F8"
GOLD = "B78024"
TEXT = "25313C"
MUTED = "667085"


def font(size: int, bold: bool = False, color: str = TEXT):
    return ImageFont.truetype(r"C:\Windows\Fonts\msyh.ttc", size=size, index=0)


def diagram(path: Path, title: str, columns: list[list[tuple[str, str]]]) -> None:
    image = Image.new("RGB", (1600, 900), "#FBFAF6")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((50, 35, 1550, 130), 24, fill=f"#{NAVY}")
    draw.text((800, 82), title, anchor="mm", font=font(36), fill="white")
    col_width = 1420 / len(columns)
    for col_i, cards in enumerate(columns):
        x0 = 90 + col_i * col_width
        x1 = x0 + col_width - 55
        y = 190
        for card_title, card_body in cards:
            h = 150
            draw.rounded_rectangle((x0, y, x1, y + h), 18, fill="white", outline="#C8D8E0", width=3)
            draw.text((x0 + 22, y + 22), card_title, font=font(25), fill=f"#{TEAL}")
            draw.multiline_text((x0 + 22, y + 67), card_body, font=font(18), fill=f"#{TEXT}", spacing=8)
            y += h + 45
        if col_i < len(columns) - 1:
            ax = x1 + 15
            ay = 445
            bx = x1 + 45
            draw.line((ax, ay, bx, ay), fill=f"#{GOLD}", width=7)
            draw.polygon([(bx, ay), (bx - 18, ay - 12), (bx - 18, ay + 12)], fill=f"#{GOLD}")
    image.save(path)


def build_diagrams() -> None:
    diagram(
        ASSET / "architecture.png",
        "系统总体架构",
        [
            [("交互层", "React 游客端\nReact 管理后台"), ("现场能力", "麦克风输入\n地图与路线")],
            [("业务服务层", "FastAPI 接口\n问答编排与记录"), ("运营闭环", "低置信问题\n反馈与游客洞察")],
            [("智能能力层", "知识检索 + DeepSeek\n百炼 ASR / TTS"), ("数字人", "LiveTalking\nWav2Lip / WebRTC")],
            [("数据与外部服务", "SQLite / 知识文档\n高德地图 / API 服务"), ("运行保障", "日志 / 状态\n一键启动与停止")],
        ],
    )
    diagram(
        ASSET / "dataflow.png",
        "一次完整问答的数据流",
        [
            [("1 语音或文本", "游客输入问题\n前端生成会话标识")],
            [("2 识别与清洗", "ASR 转写\n问题规范化")],
            [("3 知识增强", "切分与检索\nTop-K 资料片段")],
            [("4 回答生成", "模型结合资料回答\n置信度与来源")],
            [("5 数字人播报", "TTS 合成语音\nWav2Lip + WebRTC")],
        ],
    )
    diagram(
        ASSET / "operations-loop.png",
        "游客 - 知识库 - 运营闭环",
        [
            [("游客交互", "提问、路线、地图\n反馈与服务请求")],
            [("系统记录", "问答日志\n置信度与响应时间")],
            [("后台治理", "发现低置信问题\n补充或修订知识")],
            [("即时生效", "重新检索验证\n持续提升回答依据")],
        ],
    )


def set_run(run, size=11, bold=False, color=TEXT, italic=False):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure(doc: Document, running_title: str, preset: str) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6 if preset == "compact" else 8)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.33
    for name, size, before, after, color in [
        ("Title", 28, 0, 8, NAVY),
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, TEAL),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = running_title
        set_run(hp.runs[0], 8.5, False, MUTED)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("50012832  |  ")
        set_run(r, 8.5, False, MUTED)
        page_field(fp)


def cover(doc: Document, title: str, subtitle: str, version: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("灵山胜境 AI 数字人导游")
    set_run(r, 13, True, GOLD)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run(r, 28, True, NAVY)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, 14, False, TEAL)
    p.paragraph_format.space_after = Pt(56)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(version)
    set_run(r, 11, True, MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("作品编号：50012832  |  2026年7月")
    set_run(r, 10, False, MUTED)
    doc.add_page_break()


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, 10.5, True, TEAL)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(item))


def _add_steps_numbered(doc: Document, items: Iterable[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_ref])
        p_pr.append(num_pr)
        set_run(p.add_run(item))


def add_steps(doc: Document, items: Iterable[str]) -> None:
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(f"步骤 {index}  "), 10.5, True, TEAL)
        set_run(p.add_run(item), 10.5, False, TEXT)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), 9.5, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            set_run(p.add_run(str(value)), 9)
    doc.add_paragraph()


def add_image(doc: Document, image_path: Path, caption: str, width=6.65) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    p.paragraph_format.space_after = Pt(3)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    set_run(c.add_run(caption), 8.5, False, MUTED, italic=True)


def next_page(doc: Document, heading: str) -> None:
    doc.add_page_break()
    add_title(doc, heading)


def toc(doc: Document, sections: list[str]) -> None:
    add_title(doc, "目录")
    for index, section in enumerate(sections, 1):
        p = doc.add_paragraph()
        set_run(p.add_run(f"{index:02d}  {section}"), 11, index <= 3, NAVY if index <= 3 else TEXT)
        p.paragraph_format.space_after = Pt(5)


def build_manual() -> Path:
    doc = Document()
    configure(doc, "产品部署和使用手册", "compact")
    cover(doc, "产品部署和使用手册", "从解压、配置到游客端与管理后台操作", "V1.0")
    sections = ["产品简介", "软硬件环境", "提交包目录", "首次部署", "配置说明", "一键启动与停止", "游客端使用", "管理后台使用", "常见问题", "附录与验收清单"]
    toc(doc, sections)

    next_page(doc, "1. 产品简介")
    add_para(doc, "本系统面向景区游客与运营人员，将官方景区资料、智能问答、语音交互、数字人播报、路线地图与运营分析整合为一套可部署的导览产品。游客通过文字或语音提问，系统检索官方资料后生成回答，并可由数字人同步播报；运营人员可在后台维护知识、查看低置信问题和游客反馈。")
    add_bullets(doc, ["核心用户：游客、景区运营人员、知识维护人员。", "核心价值：回答有资料依据、讲解更自然、服务可持续运营。", "默认场景：Windows 单机演示，前后端与 LiveTalking 同机运行。"])
    add_image(doc, SCREEN / "01-游客端首页.png", "图1 游客端真实运行界面", 6.6)

    next_page(doc, "2. 软硬件环境")
    add_table(doc, ["类别", "建议配置", "说明"], [
        ["操作系统", "Windows 10/11 64位", "需允许 PowerShell 脚本执行"],
        ["GPU", "NVIDIA 独立显卡，建议8GB以上显存", "Wav2Lip 与实时视频处理"],
        ["CUDA/PyTorch", "与显卡驱动及 LiveTalking 环境匹配", "版本不匹配会导致模型加载失败"],
        ["Python/Conda", "Conda + livetalking 环境；后端 Python 3.10/3.11", "分别运行数字人和 FastAPI"],
        ["Node.js", "Node.js 18+ 与 npm", "运行 React/Vite 前端"],
        ["浏览器", "Chrome 或 Edge 最新版", "需允许麦克风与 WebRTC"],
        ["网络", "可访问大模型、百炼语音与高德地图服务", "离线时部分联网功能不可用"],
    ], [1.15, 2.45, 3.05])
    add_para(doc, "验收前建议执行 nvidia-smi、conda env list、python --version、node --version 和 npm --version，并确认 8010、8000、5173 端口未被其他程序占用。")

    next_page(doc, "3. 提交包目录")
    add_para(doc, "将两个 ZIP 解压到同一目录，必须形成下面的同级结构：")
    add_table(doc, ["路径", "用途"], [["启动系统.bat", "检查环境并启动三项服务"], ["停止系统.bat", "结束本次启动的服务进程"], ["使用前请阅读.txt", "一页快速说明"], ["ai-man/", "React 前端、FastAPI 后端、知识库和文档"], ["LiveTalking/", "Wav2Lip、形象数据与 WebRTC 数字人服务"]], [2.2, 4.45])
    add_para(doc, "启动脚本以自身目录为根路径，自动定位 ai-man 和 LiveTalking，不依赖 D 盘或开发者电脑上的绝对路径。")

    next_page(doc, "4. 首次部署")
    add_steps(doc, ["安装 Miniconda/Anaconda、Node.js、NVIDIA 驱动与匹配的 CUDA。", "按 LiveTalking README 创建名为 livetalking 的 Conda 环境并安装依赖。", "确认 LiveTalking/models/wav2lip.pth、data/avatars/626/coords.pkl、full_imgs 与 face_imgs 存在。", "进入 ai-man/backend，安装 requirements.txt。", "进入 ai-man/frontend，执行 npm install；源码包已保留 node_modules 时仍建议核对依赖。", "在 ai-man/backend/.env 中核对大模型、百炼语音和地图相关配置。", "双击启动系统.bat，等待浏览器打开游客端。"])
    add_table(doc, ["组件", "检查命令/文件", "成功标志"], [["LiveTalking", "conda env list", "存在 livetalking 环境"], ["Wav2Lip", "models/wav2lip.pth", "模型文件约205MB"], ["形象626", "data/avatars/626", "坐标与图像帧齐全"], ["后端", "GET /api/health", "返回 status=ok"], ["前端", "http://127.0.0.1:5173", "页面可访问"]], [1.25, 2.45, 2.95])

    next_page(doc, "5. 配置说明")
    add_table(doc, ["配置组", "主要变量", "作用"], [
        ["大模型", "LLM_PROVIDER / LLM_API_KEY / LLM_BASE_URL / LLM_MODEL", "回答生成服务"],
        ["检索", "CHAT_TOP_K / CHAT_RELIABILITY_THRESHOLD / CHAT_HISTORY_TURNS", "资料数量、可靠阈值与上下文"],
        ["语音", "BAILIAN_API_KEY / ASR_URL / TTS_URL / VOICE", "语音识别与合成"],
        ["数字人", "DIGITAL_HUMAN_BASE_URL / AVATAR / LIVETALKING_ROOT", "数字人地址、形象与目录"],
        ["数据", "DATABASE_PATH", "运行数据库位置；留空使用默认路径"],
    ], [1.2, 3.35, 2.1])
    add_para(doc, "安全提示：手册与截图不展示真实密钥。比赛包中的有效 Key 应设置调用额度和有效期，比赛结束后立即轮换或停用。")

    next_page(doc, "6. 一键启动与停止")
    add_steps(doc, ["双击启动系统.bat。", "脚本检查 ai-man 与 LiveTalking 目录、Python/Conda、前端依赖和关键端口。", "启动 LiveTalking（8010）、FastAPI（8000）和 Vite（5173）。", "浏览器自动打开游客端；启动窗口显示服务地址与日志目录。", "演示结束后双击停止系统.bat。"])
    add_table(doc, ["服务", "默认地址", "日志"], [["LiveTalking", "http://127.0.0.1:8010", "%TEMP%/ai-man-local-run/logs/livetalking.log"], ["FastAPI", "http://127.0.0.1:8000", "%TEMP%/ai-man-local-run/logs/backend.log"], ["React/Vite", "http://127.0.0.1:5173", "%TEMP%/ai-man-local-run/logs/frontend.log"]], [1.25, 2.2, 3.2])

    next_page(doc, "7. 游客端使用")
    add_image(doc, SCREEN / "01-游客端首页.png", "图2 游客端问答与数字人区域", 6.6)
    add_steps(doc, ["在问题框输入景点、演出、路线或游客服务问题，或点击麦克风进行语音提问。", "点击“开始讲解”，查看回答、资料来源与数字人播报。", "通过“重新连接/暂停播报/播报回答”控制数字人连接和语音。", "使用底部入口打开游览路线、地图导航、演出时间与游客服务。", "对回答提交游客反馈，帮助后台识别改进点。"])

    next_page(doc, "8. 管理后台使用")
    add_image(doc, SCREEN / "02-管理后台.png", "图3 管理后台总览", 6.6)
    add_bullets(doc, ["总览：查看问答量、响应时间、模型调用分布和待处理事项。", "游客洞察：汇总反馈、服务需求和人群特征。", "问答质检：查看问答记录和低置信问题，执行补知识、优化问法、标记无关或人工复核。", "知识库：新增、上传、编辑、删除和检索知识文档。", "数字人形象：查看本地形象状态、切换当前形象和管理生成任务。"])

    next_page(doc, "8.1 知识库管理")
    add_image(doc, SCREEN / "03-知识库管理.png", "图4 知识库管理真实界面", 6.6)
    add_steps(doc, ["进入“知识库”标签页。", "通过新增或上传导入官方资料，填写标题和分类。", "使用检索框验证关键问题能否命中相关原文。", "更新或删除错误/过期资料；修改后立即重新提问验证。", "对数字、票价和演出时间使用短事实卡，避免长段落混入无关内容。"])

    next_page(doc, "8.2 数字人形象管理")
    add_image(doc, SCREEN / "04-数字人形象管理.png", "图5 数字人形象管理真实界面", 6.6)
    add_steps(doc, ["进入“数字人形象”标签页，查看 coords.pkl、全身帧和人脸帧是否就绪。", "选择形象 626 作为默认演示形象；如切换形象，确认对应数据完整。", "若创建新形象任务，填写模型、形象编号和视频路径或上传视频。", "切换后回到游客端重新连接，确认画面、语音与口型正常。"])

    next_page(doc, "9. 常见问题")
    add_table(doc, ["现象", "可能原因", "处理方法"], [
        ["LiveTalking 无法启动", "环境/模型/形象缺失", "核对 livetalking 环境、wav2lip.pth 与形象626"],
        ["CUDA 报错", "驱动、CUDA 与 PyTorch 不匹配", "按显卡驱动重新安装对应 PyTorch"],
        ["端口占用", "已有进程监听8010/8000/5173", "运行停止脚本或释放端口后重试"],
        ["麦克风不可用", "浏览器权限被拒绝", "在浏览器站点设置中允许麦克风"],
        ["数字人反复重连", "LiveTalking 未就绪或会话占满", "检查8010日志并确认 max_session=2"],
        ["模型调用失败", "Key无效、余额不足或网络异常", "核对.env、额度和网络；查看后端日志"],
        ["地图失败", "Key、域名白名单或网络问题", "核对高德配置与控制台限制"],
    ], [1.55, 2.2, 2.9])

    next_page(doc, "10. 附录与快速验收")
    add_bullets(doc, ["目录：两个 ZIP 解压后五项根目录内容齐全。", "服务：8010、8000、5173 均可访问。", "问答：至少完成一题官方资料问答并显示来源。", "语音：完成一次语音识别和 TTS。", "数字人：形象626连接成功并播报回答。", "地图：路线与步行导航可打开。", "后台：知识新增后再次提问可检索生效。", "停止：停止系统.bat 能结束本次服务。"])
    add_para(doc, "本次修复后真实模型测试共20题，20次均调用 DeepSeek deepseek-v4-pro，20题均检索到官方资料；平均响应时间5.314秒，范围2.657-7.316秒。测试结果采用官方资料与数字人回复逐题对照，不设置主观正确/错误判定。")
    path = OUT / "产品部署和使用手册.docx"
    doc.save(path)
    return path


def build_design() -> Path:
    doc = Document()
    configure(doc, "产品总体设计文档", "narrative")
    cover(doc, "产品总体设计文档", "面向景区服务、知识治理与数字人交互的系统设计", "V1.0")
    sections = ["项目背景与问题分析", "需求与用户角色", "产品目标与功能边界", "总体架构", "核心数据流", "模型与知识库设计", "数字人设计", "游客端设计", "管理后台设计", "数据与接口设计", "创新点", "测试与结果", "局限与后续方向"]
    toc(doc, sections)

    next_page(doc, "1. 项目背景与问题分析")
    add_para(doc, "景区导览信息通常分散在官网、公告、地图和现场标识中，游客需要在有限时间内完成信息查找、路线决策和文化理解。固定讲解内容难以覆盖老人、亲子、半日游等差异化需求，人工服务又受人员数量和高峰客流约束。")
    add_bullets(doc, ["信息分散：景点事实、演出时间、票务与游客服务缺少统一入口。", "讲解单向：传统语音或图文无法针对追问动态调整。", "运营断层：游客问题和反馈难以直接转化为知识维护任务。", "数字人孤立：仅展示形象而缺少知识、语音、地图和后台支撑。"])

    next_page(doc, "2. 需求与用户角色")
    add_table(doc, ["角色", "核心任务", "关键需求"], [["游客", "提问、听讲解、查路线和服务", "准确、及时、自然、可追问"], ["景区运营人员", "掌握游客问题与服务热点", "看板、反馈、低置信治理"], ["知识维护人员", "更新官方资料并验证生效", "文档管理、检索验证、版本意识"], ["技术保障人员", "部署、监控与故障排查", "一键启动、日志、状态和降级"]], [1.25, 2.4, 3.0])
    add_para(doc, "系统把游客端体验与运营后台视为同一闭环的两个端点：前台产生真实问题，后台通过知识维护和质检改进下一轮回答。")

    next_page(doc, "3. 产品目标与功能边界")
    add_bullets(doc, ["提供基于官方资料的景区问答和来源展示。", "支持文本、语音、数字人播报三种交互形态。", "提供路线、地图、演出和游客服务入口。", "记录问答、置信度、模型状态和响应时间。", "支持知识库、低置信问题、游客洞察和数字人形象管理。"])
    add_para(doc, "功能边界：系统不替代景区现场安全指挥、医疗救助或官方票务交易；实时演出、票价与开放政策必须以景区当日公告为准。")

    next_page(doc, "4. 总体架构")
    add_image(doc, ASSET / "architecture.png", "图1 系统总体架构", 6.6)
    add_para(doc, "架构采用前后端分离与能力解耦设计。React 负责游客端和管理后台；FastAPI 统一编排问答、知识、语音、数字人和运营数据；DeepSeek 与百炼提供联网智能能力；LiveTalking 负责基于 Wav2Lip 的形象驱动与 WebRTC 输出；SQLite 和知识文档承担本地数据存储。")

    next_page(doc, "4.1 组件职责")
    add_table(doc, ["组件", "职责", "设计理由"], [["React/Vite", "交互、状态与可视化", "开发轻量，适合单机演示和快速迭代"], ["FastAPI", "API与业务编排", "异步接口适合模型、语音和数字人调用"], ["知识检索", "从官方资料选取相关片段", "降低无依据生成，支持来源展示"], ["DeepSeek", "结合资料生成自然回答", "兼顾中文表达与联网模型能力"], ["百炼语音", "ASR和TTS", "实时语音链路与中文音色"], ["LiveTalking", "Wav2Lip与WebRTC", "将回答转换为可见可听的数字人讲解"], ["SQLite", "问答、反馈、知识和状态", "单机部署简单、无需独立数据库服务"]], [1.15, 2.25, 3.25])

    next_page(doc, "5. 核心数据流")
    add_image(doc, ASSET / "dataflow.png", "图2 一次问答的数据流", 6.6)
    add_steps(doc, ["游客输入文本或语音；语音先经 ASR 转写。", "后端清洗问题并结合会话历史抽取检索意图。", "知识服务切分、评分并返回 Top-K 官方资料。", "Prompt 约束模型仅依据资料回答，并输出来源、置信度和可靠标记。", "回答由 TTS 合成语音，LiveTalking 驱动形象并通过 WebRTC 播放。", "问答、耗时和模型状态进入运营数据，低置信问题形成治理任务。"])

    next_page(doc, "6. 模型与知识库设计")
    add_para(doc, "知识库设计重点不在堆叠长文，而在让每个问题能稳定检索到完整、单一主题的官方事实。资料通过段落与滑动窗口构建检索块；对票价、尺寸、寓意等易被长文淹没的事实，增加短事实卡并禁用跨主题滑动窗口。")
    add_bullets(doc, ["切分：保留段落语义，同时提供必要上下文窗口。", "排序：关键词、实体、问题类别和负向规则共同决定相关度。", "Prompt：要求依据资料回答，缺少依据时明确说明，不编造实时信息。", "置信度：低于阈值时标记为不可靠并进入后台治理。", "来源：API返回检索片段，游客端和测试结果可直接对照。", "上下文：保留有限轮次，支持追问但避免历史污染。"])

    next_page(doc, "6.1 官方事实冲突治理")
    add_para(doc, "官方资料包内部可能出现数字冲突。系统对九龙灌浴高度等关键事实采用单一权威口径，并将冲突值从默认知识文本中清理。修复后增加登云道寓意、九龙灌浴尺寸、梵宫尺寸、祥符禅寺遗存和票价等短事实卡，确保完整事实进入前三条检索结果。")
    add_table(doc, ["事实主题", "统一口径/处理"], [["登云道216级", "前108级烦恼尽除，后108级愿望圆满"], ["九龙灌浴", "总高27.5米，中央太子佛像7.2米"], ["灵山梵宫", "建筑面积72000平方米，最高处66.5米"], ["票价", "成人210元、半价105元、观光车40元/人；以当日公告为准"], ["事实分块", "短卡不生成跨主题滑动窗口，避免无关内容混入"]], [1.6, 5.05])

    next_page(doc, "7. 数字人设计")
    add_para(doc, "数字人链路由文本回答、TTS语音、Wav2Lip口型驱动、形象帧数据和WebRTC传输组成。默认使用形象626，并保留形象708及其他模型素材。数字人服务与问答服务分离，LiveTalking不可用时文本问答、资料来源和后台仍可继续工作。")
    add_bullets(doc, ["模型：Wav2Lip，默认检查 models/wav2lip.pth。", "形象：data/avatars/626 与 708，包含坐标、全身帧和人脸帧。", "传输：浏览器通过 /offer 建立 WebRTC，会话数默认2。", "播报：后端回答完成后调用 LiveTalking /human。", "故障降级：显示重连状态，保留文本回答与人工播报按钮。"])

    next_page(doc, "8. 游客端设计")
    add_image(doc, SCREEN / "01-游客端首页.png", "图3 游客端真实界面", 6.6)
    add_para(doc, "游客端采用“左侧数字人、右侧问答、底部服务入口”的布局。问题示例降低首次使用门槛，语音按钮突出自然交互；数字人连接状态、重新连接、暂停与播报按钮使链路状态可见、可控。")

    next_page(doc, "9. 管理后台设计")
    add_image(doc, SCREEN / "02-管理后台.png", "图4 管理后台真实界面", 6.6)
    add_para(doc, "后台围绕运营闭环而非单纯日志展示：总览呈现问答量、响应时间和模型分布；低置信问题提供补充知识、优化问法、标记无关和人工复核动作；游客洞察、知识库和形象管理共同支撑持续运营。")

    next_page(doc, "9.1 知识治理闭环")
    add_image(doc, ASSET / "operations-loop.png", "图5 游客-知识库-运营闭环", 6.6)
    add_bullets(doc, ["真实问题进入问答记录。", "置信度和资料来源暴露潜在缺口。", "运营人员在知识库新增或修订官方内容。", "检索与问答即时使用新知识。", "游客反馈和后续问题验证改进效果。"])

    next_page(doc, "10. 数据与接口设计")
    add_table(doc, ["接口组", "代表接口", "用途"], [["健康", "GET /api/health", "启动与监控检查"], ["问答", "POST /api/chat", "清洗、检索、模型回答与记录"], ["语音", "POST /api/speech/recognize、synthesize", "ASR与TTS"], ["知识", "GET/POST/PUT/DELETE /api/admin/knowledge/*", "文档管理与检索"], ["问答记录", "GET /api/admin/chat-records、dashboard", "质检与运营看板"], ["数字人", "GET /api/digital-human/config、avatars", "连接配置与形象管理"], ["反馈", "POST /api/feedback", "游客评价与帮助率"], ["游客报告", "GET /api/admin/visitor-report", "游客行为数据"]], [1.1, 3.15, 2.4])
    add_para(doc, "运行数据默认存储在 data/runtime 下，包括 SQLite 数据库与数字人选择状态；启动脚本把日志和 PID 文件写入系统临时目录，避免提交目录被频繁写入。")

    next_page(doc, "11. 创新点")
    add_bullets(doc, ["景区专属知识增强：以官方资料为依据，并显示检索来源。", "多模态数字人输出：打通 ASR、LLM、TTS、Wav2Lip 和 WebRTC。", "游客-知识库-运营闭环：真实问题直接驱动知识维护。", "低置信治理：将不确定回答转化为可处理任务，而非隐藏模型风险。", "可扩展架构：模型、语音、数字人、地图和知识模块通过接口解耦。"])

    next_page(doc, "12. 测试与结果")
    add_para(doc, "知识库修复采用回归测试验证关键事实的前三条检索结果，并对知识、问答和数字人相关用例执行自动化测试。真实模型测试按用户要求不设主观正确/错误标签，而是并列展示官方资料相关内容、数字人实际回复和响应时间。")
    add_table(doc, ["指标", "结果"], [["测试问题数", "20"], ["真实模型完成数", "20"], ["模型", "DeepSeek deepseek-v4-pro"], ["检索到官方资料的问题", "20/20，每题3条来源"], ["平均响应时间", "5.314秒"], ["响应时间范围", "2.657-7.316秒"], ["自动化相关回归", "知识库/问答27项通过；路径回归5项通过"]], [2.4, 4.25])
    add_para(doc, "详细逐题结果见《灵山胜境数字人问答对照测试结果-修复后.xlsx》。")

    next_page(doc, "13. 局限与后续方向")
    add_bullets(doc, ["联网依赖：大模型、语音和地图功能受网络与第三方服务可用性影响。", "GPU要求：实时数字人需要合适的NVIDIA显卡、CUDA与PyTorch环境。", "资料时效：演出、票价和开放政策需持续同步官方公告。", "单机规模：当前适合比赛演示和单点部署，规模化需引入独立数据库、任务队列和统一监控。", "多景区扩展：后续可增加租户/景区隔离、知识版本、权限与端侧轻量部署。"])
    add_para(doc, "设计结论：本作品不是单一聊天页面，而是一套可部署、可讲解、可维护、可观察并能持续改进的景区数字人导览系统。")
    path = OUT / "产品总体设计文档.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    build_diagrams()
    print(build_manual())
    print(build_design())
