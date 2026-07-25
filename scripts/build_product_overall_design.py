from __future__ import annotations

import re
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "tmp" / "product-overall-design"
CONTENT = WORK / "draft_content.md"
FIGURES = WORK / "figures"
OUTPUT = ROOT / "outputs" / "灵山胜境景区导览AI数字人-产品总体设计文档.docx"

TEAL = "155E5A"
DEEP_TEAL = "0E4644"
GOLD = "B88A35"
INK = "111111"
MUTED = "5E6B68"
PALE = "F4F7F5"
PALE_GOLD = "FBF7ED"
WHITE = "FFFFFF"
GRID = "AEB8B5"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = Path(r"C:\Windows\Fonts\msyh.ttc")
    return ImageFont.truetype(str(path), size=size, index=0)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt: ImageFont.FreeTypeFont) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        if current and draw.textbbox((0, 0), trial, font=fnt)[2] > max_width:
            lines.append(current)
            current = char
        else:
            current = trial
    if current:
        lines.append(current)
    return lines


def rounded_card(draw, box, title, body, fill=WHITE, accent=TEAL, title_size=27, body_size=20):
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=f"#{fill}", outline=f"#{GRID}", width=2)
    draw.rounded_rectangle((x0, y0, x0 + 12, y1), radius=6, fill=f"#{accent}")
    draw.text((x0 + 28, y0 + 22), title, font=font(title_size, True), fill=f"#{DEEP_TEAL}")
    y = y0 + 68
    body_font = font(body_size)
    for line in wrap_text(draw, body, x1 - x0 - 56, body_font):
        draw.text((x0 + 28, y), line, font=body_font, fill=f"#{INK}")
        y += body_size + 13


def canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((800, 60), title, anchor="mm", font=font(36, True), fill=f"#{DEEP_TEAL}")
    draw.line((120, 105, 1480, 105), fill=f"#{GOLD}", width=4)
    return image, draw


def arrow(draw, start, end, color=GOLD, width=6):
    x0, y0 = start
    x1, y1 = end
    draw.line((x0, y0, x1, y1), fill=f"#{color}", width=width)
    import math

    angle = math.atan2(y1 - y0, x1 - x0)
    size = 18
    p1 = (x1 - size * math.cos(angle - 0.55), y1 - size * math.sin(angle - 0.55))
    p2 = (x1 - size * math.cos(angle + 0.55), y1 - size * math.sin(angle + 0.55))
    draw.polygon([(x1, y1), p1, p2], fill=f"#{color}")


def build_value_map(path: Path):
    image, draw = canvas("从景区问题到可验证价值")
    headings = [("现有问题", 110), ("产品能力", 585), ("产生价值", 1060)]
    for title, x in headings:
        draw.rounded_rectangle((x, 135, x + 420, 205), radius=15, fill=f"#{DEEP_TEAL}")
        draw.text((x + 210, 170), title, anchor="mm", font=font(27, True), fill="white")
    rows = [
        ("信息入口分散", "问答与游览服务统一入口", "游客少查找、少切换"),
        ("模型回答依据不清", "景区检索、来源和可靠状态", "事实可追溯，风险可见"),
        ("数字人停留在展示", "语音、问答和实时播报协同", "讲解自然且能承接追问"),
        ("问题与反馈不沉淀", "后台质检、知识维护和复测", "真实使用推动持续改进"),
    ]
    y = 235
    for left, mid, right in rows:
        rounded_card(draw, (110, y, 530, y + 120), left, "", fill=PALE_GOLD, accent=GOLD, title_size=23)
        rounded_card(draw, (585, y, 1005, y + 120), mid, "", fill=PALE, accent=TEAL, title_size=23)
        rounded_card(draw, (1060, y, 1480, y + 120), right, "", fill=PALE, accent=TEAL, title_size=23)
        arrow(draw, (535, y + 60), (575, y + 60))
        arrow(draw, (1010, y + 60), (1050, y + 60))
        y += 145
    image.save(path)


def build_overall_solution(path: Path):
    image, draw = canvas("游客服务与景区管理共用同一能力底座")
    rounded_card(draw, (110, 145, 690, 300), "游客端", "文字与语音提问｜数字人讲解｜路线、地图、演出与游客服务", fill=PALE, accent=TEAL)
    rounded_card(draw, (910, 145, 1490, 300), "管理后台", "运营总览｜游客洞察｜问答质检｜知识与数字人管理", fill=PALE_GOLD, accent=GOLD)
    rounded_card(draw, (420, 365, 1180, 535), "FastAPI 业务与数据中枢", "统一编排问答、知识、记录、反馈、语音、数字人和地图接口", fill="EEF5F2", accent=DEEP_TEAL, title_size=29, body_size=22)
    arrow(draw, (400, 300), (600, 355))
    arrow(draw, (1200, 300), (1000, 355))
    items = [
        ("景区知识与检索", "资料切分、匹配、来源"),
        ("大语言模型", "基于资料组织导览回答"),
        ("语音与数字人", "ASR、TTS、Wav2Lip、WebRTC"),
        ("地图与运行数据", "步行导航、SQLite、状态记录"),
    ]
    x = 80
    for title, body in items:
        rounded_card(draw, (x, 625, x + 350, 790), title, body, fill=WHITE, accent=TEAL, title_size=22, body_size=17)
        arrow(draw, (800, 545), (x + 175, 615), color=GRID, width=4)
        x += 390
    image.save(path)


def build_runtime_architecture(path: Path):
    image, draw = canvas("前后端分离、能力接口化的运行架构")
    layers = [
        ("交互层", "React 游客端 ｜ React 管理后台", PALE_GOLD, GOLD),
        ("业务层", "FastAPI 接口编排 ｜ 会话与问答处理 ｜ 反馈与运营分析", PALE, TEAL),
        ("智能能力层", "景区知识检索 ｜ 兼容 OpenAI 接口的大语言模型 ｜ ASR / TTS", "EEF5F2", DEEP_TEAL),
        ("数字人与地图", "LiveTalking ｜ Wav2Lip ｜ WebRTC ｜ 高德步行导航", PALE, TEAL),
        ("数据层", "knowledge.db ｜ chat_records.db ｜ 景区资料与数字人状态", PALE_GOLD, GOLD),
    ]
    y = 135
    for title, body, fill, accent in layers:
        rounded_card(draw, (170, y, 1430, y + 112), title, body, fill=fill, accent=accent, title_size=25, body_size=20)
        y += 135
    image.save(path)


def build_operations_loop(path: Path):
    image, draw = canvas("游客问题驱动景区内容持续更新")
    nodes = [
        ((160, 180, 650, 340), "游客真实交互", "提问、追问、评价与服务请求"),
        ((950, 180, 1440, 340), "结构化记录", "问题、来源、置信度、状态与耗时"),
        ((950, 560, 1440, 720), "管理员处理", "补知识、改内容、标无关或人工复核"),
        ((160, 560, 650, 720), "复测并生效", "用原问题验证，新内容进入后续问答"),
    ]
    for box, title, body in nodes:
        rounded_card(draw, box, title, body, fill=PALE, accent=TEAL, title_size=27, body_size=20)
    arrow(draw, (660, 260), (940, 260))
    arrow(draw, (1195, 350), (1195, 550))
    arrow(draw, (940, 640), (660, 640))
    arrow(draw, (405, 550), (405, 350))
    draw.ellipse((650, 360, 950, 540), fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=4)
    draw.text((800, 435), "知识与服务", anchor="mm", font=font(29, True), fill=f"#{DEEP_TEAL}")
    draw.text((800, 480), "随真实使用改进", anchor="mm", font=font(22), fill=f"#{MUTED}")
    image.save(path)


def build_answer_flow(path: Path):
    image, draw = canvas("一次有依据问答的处理过程")
    steps = [
        ("1", "游客提问", "文字或语音"),
        ("2", "问题处理", "清洗与上下文"),
        ("3", "知识检索", "匹配景区资料"),
        ("4", "回答生成", "模型依据资料表达"),
        ("5", "可靠判断", "来源、置信度与状态"),
        ("6", "展示播报", "文字、语音与数字人"),
        ("7", "记录反馈", "进入后台治理"),
    ]
    x = 55
    y = 260
    width = 190
    for index, title, body in steps:
        draw.ellipse((x + 63, y - 70, x + 127, y - 6), fill=f"#{DEEP_TEAL}")
        draw.text((x + 95, y - 38), index, anchor="mm", font=font(25, True), fill="white")
        rounded_card(draw, (x, y, x + width, y + 260), title, body, fill=PALE, accent=TEAL, title_size=22, body_size=18)
        if index != "7":
            arrow(draw, (x + width + 5, y + 130), (x + width + 25, y + 130), width=4)
        x += 220
    draw.rounded_rectangle((250, 660, 1350, 775), radius=20, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=2)
    draw.text((800, 705), "资料不足时明确提示并进入低置信度治理；数字人异常时保留文本链路", anchor="mm", font=font(23, True), fill=f"#{DEEP_TEAL}")
    image.save(path)


def build_figures():
    FIGURES.mkdir(parents=True, exist_ok=True)
    build_value_map(FIGURES / "value-map.png")
    build_overall_solution(FIGURES / "overall-solution.png")
    build_runtime_architecture(FIGURES / "runtime-architecture.png")
    build_operations_loop(FIGURES / "operations-loop.png")
    build_answer_flow(FIGURES / "answer-flow.png")


def set_east_asia_font(run, east_asia="宋体", latin="Times New Roman"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def set_run(run, size=12, bold=False, italic=False, color=INK, east_asia="宋体"):
    set_east_asia_font(run, east_asia=east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    specs = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (14, 14, 8),
        "Heading 3": (12.5, 10, 6),
    }
    for name, (size, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    settings = doc.settings.element
    if settings.find(qn("w:doNotAutoCompressPictures")) is None:
        settings.append(OxmlElement("w:doNotAutoCompressPictures"))


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=20, bold=True, east_asia="微软雅黑")


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    match = re.match(r"^\*\*(.+?)\*\*\s*(.*)$", text)
    if match:
        r = p.add_run(match.group(1))
        set_run(r, bold=True, east_asia="微软雅黑")
        if match.group(2):
            r = p.add_run(" " + match.group(2))
            set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E9EFEC")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=10.5, bold=True, east_asia="微软雅黑")
    prevent_row_split(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, width=5.77):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Pt(0)
    set_run(c.add_run(caption), size=10, bold=True, east_asia="微软雅黑")


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    def cells(line: str) -> list[str]:
        return [item.strip() for item in line.strip().strip("|").split("|")]

    headers = cells(lines[start])
    index = start + 2
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append(cells(lines[index]))
        index += 1
    return headers, rows, index


def build_document():
    doc = Document()
    configure_document(doc)
    lines = CONTENT.read_text(encoding="utf-8").splitlines()
    first_title = True
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line == "[[PAGEBREAK]]":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue
        figure = re.match(r"^\[(FIGURE|SCREENSHOT):(.+?)\|(.+?)\]$", line)
        if figure:
            kind, path_text, caption = figure.groups()
            image_path = FIGURES / path_text if kind == "FIGURE" else Path(path_text)
            add_figure(doc, image_path, caption)
            index += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if first_title:
                add_title(doc, text)
                first_title = False
            else:
                doc.add_heading(text, level=1)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1].strip()):
            headers, rows, next_index = parse_table(lines, index)
            add_table(doc, headers, rows)
            index = next_index
            continue
        add_body(doc, line)
        index += 1

    doc.core_properties.title = "灵山胜境景区导览 AI 数字人产品总体设计文档"
    doc.core_properties.subject = "需求场景、整体方案、核心机制、技术支撑与项目优势"
    doc.core_properties.author = "灵山胜境景区导览 AI 数字人项目组"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    build_figures()
    print(build_document())
